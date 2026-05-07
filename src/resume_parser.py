"""Resume text extraction from PDF, DOCX, and TXT files."""

import io
import streamlit as st


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyPDF2."""
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required: pip install PyPDF2")

    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
            pages.append(text)
        except Exception:
            pages.append("")
    return "\n".join(pages)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes including tables."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def parse_resume(uploaded_file) -> str | None:
    """Parse uploaded resume file and return extracted text."""
    if uploaded_file is None:
        return None

    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    try:
        if name.endswith(".pdf"):
            return parse_pdf(file_bytes)
        elif name.endswith(".docx"):
            return parse_docx(file_bytes)
        elif name.endswith(".txt"):
            return file_bytes.decode("utf-8", errors="replace")
        else:
            st.error("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")
            return None
    except Exception as exc:
        st.error(f"Failed to parse resume: {exc}")
        return None
